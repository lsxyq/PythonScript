#!/usr/bin/env python
# -*- coding: UTF-8 -*-
# Author:Leslie-x 

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm


# tb =doc.tables

# rows = tb[0].rows  #取出第一张表的所有行
# cols = rows[0].cells #取出第一行的所有列
# cell = cols[0] #取出第一行第一列的单元格


# print "row,col,row"
# print df.shape[0],df.shape[1],len(df)


def excel_to_doc(excel_path, doc_path):
    doc = Document()
    df = pd.read_excel(excel_path, sheetname='Sheet1')
    headers = list(df.keys())[1:]
    info = {}
    for header in headers:
        info[header] = list(df[header])
    tb = doc.add_table(rows=len(df), cols=len(headers))
    tb.add_row()

    for i in range(len(headers)):
        tb.cell(0, i).text = headers[i]

    for row in range(1, len(df) + 1):
        index = row - 1
        for col, header in enumerate(headers):
            tb.cell(row, col).width = 1
            tb.cell(row, col).text = str(info.get(header)[index])
            tb.cell(row, col).width = Cm(6)
    tb.style = 'Medium Grid 1 Accent 1'
    doc.save(doc_path)


excel_path = 'F:\website.xlsx'
doc_path = 'website.docx'
excel_to_doc(excel_path, doc_path)
